import hashlib
import io
import json
import sqlite3
from datetime import datetime
from pathlib import Path

from fastapi import HTTPException
from docx import Document

from app.core.artifact_crypto import read_artifact_bytes, write_artifact_bytes
from app.core.settings_store import get_settings
from app.db.database import get_connection
from app.services.analysis import run_analysis
from app.services.audit import write_audit_event


def _get_resume_and_posting(resume_version_id: int, job_posting_id: int) -> tuple[tuple, tuple]:
    with get_connection() as conn:
        resume = conn.execute(
            """
            SELECT id, source_name, file_path
            FROM resume_versions
            WHERE id = ?
            """,
            (resume_version_id,),
        ).fetchone()

        posting = conn.execute(
            """
            SELECT id, title
            FROM job_postings
            WHERE id = ?
            """,
            (job_posting_id,),
        ).fetchone()

    if not resume:
        raise HTTPException(status_code=404, detail={"code": "NOT_FOUND", "message": "Resume version not found."})
    if not posting:
        raise HTTPException(status_code=404, detail={"code": "NOT_FOUND", "message": "Job posting not found."})

    return resume, posting


def _next_version_tag(conn: sqlite3.Connection, source_name: str) -> str:
    row = conn.execute("SELECT COUNT(*) FROM resume_versions WHERE source_name = ?", (source_name,)).fetchone()
    next_version = (int(row[0]) if row else 0) + 1
    return f"v{next_version}"


def run_optimization(resume_version_id: int, job_posting_id: int) -> dict:
    try:
        settings = get_settings()
        resume, posting = _get_resume_and_posting(resume_version_id, job_posting_id)
        analysis = run_analysis(resume_version_id, job_posting_id)

        source_name = resume[1]
        source_path = Path(resume[2])
        if not source_path.exists():
            raise HTTPException(status_code=400, detail={"code": "VALIDATION_ERROR", "message": "Source resume file missing."})

        source_bytes = read_artifact_bytes(source_path)

        timestamp = datetime.now().strftime("%Y%m%d%H%M")
        deterministic_name = f"{source_name}__opt__job{job_posting_id}__{timestamp}.docx"
        target_path = source_path.parent / deterministic_name
        counter = 1
        while target_path.exists() or target_path.with_suffix(target_path.suffix + ".enc").exists():
            deterministic_name = f"{source_name}__opt__job{job_posting_id}__{timestamp}__r{counter}.docx"
            target_path = source_path.parent / deterministic_name
            counter += 1

        document = Document(io.BytesIO(source_bytes))
        document.add_paragraph("")
        document.add_heading("Optimization Notes", level=2)
        document.add_paragraph(f"Target Role: {posting[1]}")

        if analysis.missing_keywords.hard:
            document.add_paragraph("Hard skills to consider adding:")
            for skill in analysis.missing_keywords.hard:
                document.add_paragraph(f"- {skill}")

        if analysis.missing_keywords.soft:
            document.add_paragraph("Soft skills to emphasize:")
            for skill in analysis.missing_keywords.soft:
                document.add_paragraph(f"- {skill}")

        document.add_paragraph(
            "Review and manually tailor these suggestions to your real experience before submitting applications."
        )

        output_buffer = io.BytesIO()
        document.save(output_buffer)
        body = output_buffer.getvalue()
        stored_target_path = write_artifact_bytes(target_path, body, encrypt=settings.storage.encrypt_artifacts)
        checksum = hashlib.sha256(body).hexdigest()

        suggestion_summary = {
            "overall_score": analysis.overall_score,
            "missing_hard_skills": analysis.missing_keywords.hard,
            "missing_soft_skills": analysis.missing_keywords.soft,
        }

        with get_connection() as conn:
            version_tag = _next_version_tag(conn, source_name)
            conn.execute(
                """
                INSERT INTO resume_versions (
                  source_name, version_tag, mime_type, file_ext, file_path,
                  checksum_sha256, extracted_text, sections_json, parser_confidence, parent_resume_version_id
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    source_name,
                    version_tag,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ".docx",
                    str(stored_target_path),
                    checksum,
                    "",
                    "{}",
                    0.75,
                    resume_version_id,
                ),
            )
            output_resume_version_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]

            conn.execute(
                """
                INSERT INTO optimized_artifacts (
                  source_resume_version_id,
                  output_resume_version_id,
                  output_file_path,
                  suggestion_summary_json,
                  deterministic_name
                ) VALUES (?, ?, ?, ?, ?)
                """,
                (
                    resume_version_id,
                    output_resume_version_id,
                    str(stored_target_path),
                    json.dumps(suggestion_summary),
                    deterministic_name,
                ),
            )
            conn.commit()

        write_audit_event(
            event_type="optimization.run",
            entity_type="optimized_artifact",
            entity_id=str(output_resume_version_id),
            payload={
                "source_resume_version_id": resume_version_id,
                "output_resume_version_id": output_resume_version_id,
                "output_file_path": str(stored_target_path),
                "deterministic_name": deterministic_name,
            },
        )

        return {
            "source_resume_version_id": resume_version_id,
            "output_resume_version_id": output_resume_version_id,
            "output_file_path": str(stored_target_path),
            "deterministic_name": deterministic_name,
            "suggestion_summary": suggestion_summary,
        }
    except sqlite3.OperationalError as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "code": "VALIDATION_ERROR",
                "message": "Database is not initialized. Run /api/v1/db/init first.",
                "details": {"reason": str(exc)},
            },
        )


def build_cover_letter_prompt(resume_version_id: int, job_posting_id: int, tone: str, target_length_words: int) -> dict:
    try:
        with get_connection() as conn:
            resume = conn.execute(
                """
                SELECT source_name, extracted_text
                FROM resume_versions
                WHERE id = ?
                """,
                (resume_version_id,),
            ).fetchone()
            posting = conn.execute(
                """
                SELECT c.name, jp.title, jp.description_text
                FROM job_postings jp
                JOIN companies c ON c.id = jp.company_id
                WHERE jp.id = ?
                """,
                (job_posting_id,),
            ).fetchone()
    except sqlite3.OperationalError as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "code": "VALIDATION_ERROR",
                "message": "Database is not initialized. Run /api/v1/db/init first.",
                "details": {"reason": str(exc)},
            },
        )

    if not resume:
        raise HTTPException(status_code=404, detail={"code": "NOT_FOUND", "message": "Resume version not found."})
    if not posting:
        raise HTTPException(status_code=404, detail={"code": "NOT_FOUND", "message": "Job posting not found."})

    analysis = run_analysis(resume_version_id, job_posting_id)

    strengths = analysis.matched_keywords.hard[:5] + analysis.matched_keywords.soft[:3]
    gaps = analysis.missing_keywords.hard[:5] + analysis.missing_keywords.soft[:3]

    prompt = (
        "Write a tailored cover letter using the following constraints:\n"
        f"- Company: {posting[0]}\n"
        f"- Role: {posting[1]}\n"
        f"- Tone: {tone}\n"
        f"- Target length: {target_length_words} words\n"
        f"- Candidate resume profile: {resume[0]}\n"
        f"- Emphasize strengths: {', '.join(strengths) if strengths else 'relevant accomplishments'}\n"
        f"- Address likely gaps tactfully: {', '.join(gaps) if gaps else 'none flagged'}\n"
        "- Keep claims factual and grounded in prior experience.\n"
        "- Include a concise opening, 1-2 evidence paragraphs, and a closing call to action.\n"
    )

    return {
        "prompt": prompt,
        "metadata": {
            "company": posting[0],
            "role": posting[1],
            "tone": tone,
            "target_length_words": target_length_words,
        },
    }
